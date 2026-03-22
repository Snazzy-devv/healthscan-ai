import streamlit as st
import pdfplumber
import easyocr
import numpy as np
import io
from PIL import Image
from docx import Document
import os
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, SystemMessage

# ----------------------
# Load environment variables
# ----------------------
load_dotenv(override=True)

# ----------------------
# Ensure OPENROUTER_API_KEY is set
# ----------------------
openrouter_key = os.getenv("OPENROUTER_API_KEY")
if not openrouter_key:
    st.error("❌ OPENROUTER_API_KEY not set! Add it to .env or Streamlit Secrets.")
    st.stop()
os.environ["OPENROUTER_API_KEY"] = openrouter_key

# ----------------------
# Streamlit Page Config
# ----------------------
st.set_page_config(page_title="HealthScan AI", page_icon="💊", layout="centered")

# ----------------------
# Initialize OCR reader in session state
# ----------------------
if 'ocr_reader' not in st.session_state:
    st.session_state.ocr_reader = easyocr.Reader(['en'])

# Initialize LLM in session state
if 'llm' not in st.session_state:
    st.session_state.llm = ChatOpenAI(
        model="gpt-4o",
        api_key=openrouter_key,
        base_url="https://openrouter.ai/api/v1"
    )

# ----------------------
# Helper Functions
# ----------------------
def extract_text(uploaded_file):
    """Extract text from PDF or image."""
    text = ""

    if "pdf" in uploaded_file.type:
        with pdfplumber.open(uploaded_file) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
    else:
        image = Image.open(uploaded_file)
        try:
            img_np = np.array(image)
            results = st.session_state.ocr_reader.readtext(img_np, detail=0)
            text = " ".join(results)
        finally:
            image.close()
    return text

def create_docx(markdown_content):
    """Convert Markdown to Word DOCX."""
    doc = Document()
    doc.add_heading('Medical Lab Analysis Report', 0)

    for line in markdown_content.split('\n'):
        if line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=1)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=2)
        elif line.startswith('* '):
            doc.add_paragraph(line.replace('* ', ''), style='List Bullet')
        elif line.strip():
            doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ----------------------
# Streamlit UI
# ----------------------
st.title("🩺 HealthScan AI")
st.subheader("Instant Lab Result Analysis & Document Converter")
st.markdown("---")

# File uploader
uploaded_file = st.file_uploader(
    "Upload a photo or PDF of your lab results",
    type=["pdf", "png", "jpg", "jpeg"],
    help="Max 10MB recommended"
)

if uploaded_file:
    # Clear previous session data
    for key in ['raw_text', 'report_md']:
        if key in st.session_state:
            del st.session_state[key]

    with st.status("Analyzing your document...", expanded=True) as status:
        st.write("📄 Reading text from file...")
        st.session_state.raw_text = extract_text(uploaded_file)

        if not st.session_state.raw_text.strip():
            st.error("❌ Could not extract text from the file.")
            st.stop()

        st.write("🤖 Consulting AI Medical Knowledge Base...")

        system_prompt = SystemMessage(content="""
You are a Medical Analyst. You receive raw text from lab scans.
Analyze the biomarkers and generate a report in Markdown.

Structure:
## 📋 Summary
## 🧪 Lab Results Table (Test, Result, Range, Status)
## ⚠️ Detected Ailments/Issues
## 💡 Solutions & Advice
- **Medications/Supplements:** (General advice)
- **Avoid:** (Foods/Activities)
- **Improvement:** (Lifestyle changes)

DISCLAIMER: This is not a clinical diagnosis.
        """)

        user_prompt = HumanMessage(content=f"Analyze this content:\n\n{st.session_state.raw_text}")

        try:
            response = st.session_state.llm.invoke([system_prompt, user_prompt])
            st.session_state.report_md = response.content
        except Exception as e:
            st.error(f"❌ AI Error: {e}")
            st.stop()

        status.update(label="✅ Analysis Complete!", state="complete", expanded=False)

    # Display AI report
    st.markdown(st.session_state.report_md)
    st.markdown("---")
    st.subheader("📥 Download Your Report")

    col1, col2 = st.columns(2)

    with col1:
        docx_data = create_docx(st.session_state.report_md)
        st.download_button(
            label="Download as Word (.docx)",
            data=docx_data,
            file_name="Medical_Analysis.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        docx_data.close()  # free memory

    with col2:
        st.download_button(
            label="Download as Markdown (.md)",
            data=st.session_state.report_md,
            file_name="Medical_Analysis.md",
            mime="text/markdown"
        )

else:
    st.info("📤 Please upload a file to begin.")
